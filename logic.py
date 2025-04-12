import os
import json
import time
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import datetime
from pathlib import Path
from google.generativeai import caching
import cv2
from typing import Dict, List
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import io

class HomeInspector:
    def __init__(self, api_key: str, standards_dir: str, examples_dir: str):
        self.api_key = api_key
        self.standards_dir = Path(standards_dir)
        self.examples_dir = Path(examples_dir)
        self.document_dict = {
            'building_standards': {},
            'examples': {'example1': {}, 'example2': {}},
            'user_data': {}
        }
        self.SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.doc', '.docx', '.jpg', '.jpeg', '.png'}
        self._initialize_model()

    def _initialize_model(self):
        genai.configure(api_key=self.api_key)
        
        # Load building standards
        self._load_standards()
        
        # Load examples
        self._load_examples()
        
        # Initialize model with cache
        self.cache = caching.CachedContent.create(
            model='models/gemini-1.5-flash-002',
            display_name='home_inspection_cache',
            system_instruction=(
                'You are an expert at analysing residential building and producing detailed inspection reports.'
                'Your job is to analyse the user provided media and produce a detailed inspection report based on the reference standards you have access to.'
            ),
            contents=[doc for doc in self.document_dict['building_standards'].values()],
            ttl=datetime.timedelta(minutes=60),
        )
        
        generation_config = {
            "temperature": 0.1,
            "max_output_tokens": 8192,
            "response_mime_type": "application/json",
        }
        
        self.model = genai.GenerativeModel.from_cached_content(
            cached_content=self.cache, 
            generation_config=generation_config
        )

    def _load_standards(self):
        for file_path in self.standards_dir.rglob('*'):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    uploaded_file = genai.upload_file(str(file_path))
                    self.document_dict['building_standards'][file_path.name] = uploaded_file
                except Exception as e:
                    print(f"Error loading standard {file_path.name}: {str(e)}")

    def _load_examples(self):
        for file_path in self.examples_dir.rglob('*'):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    subfolder = file_path.parent.name
                    if subfolder in ['example1', 'example2']:
                        uploaded_file = genai.upload_file(str(file_path))
                        self.document_dict['examples'][subfolder][file_path.name] = uploaded_file
                except Exception as e:
                    print(f"Error loading example {file_path.name}: {str(e)}")

    def process_video(self, video_path: str, output_dir: str = "extracted_frames") -> Dict[str, str]:
        """Process video and extract frames at regular intervals"""
        Path(output_dir).mkdir(exist_ok=True)
        frame_paths = {}
        
        cap = cv2.VideoCapture(video_path)
        if not cap.isOpened():
            raise ValueError(f"Could not open video file: {video_path}")
        
        fps = cap.get(cv2.CAP_PROP_FPS)
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        duration = frame_count / fps
        
        # Extract frames every 5 seconds
        for timestamp in range(0, int(duration), 5):
            cap.set(cv2.CAP_PROP_POS_MSEC, timestamp * 1000)
            ret, frame = cap.read()
            if ret:
                frame_filename = f"frame_{timestamp}.jpg"
                frame_path = str(Path(output_dir) / frame_filename)
                cv2.imwrite(frame_path, frame)
                frame_paths[f"video_{timestamp}s"] = frame_path
        
        cap.release()
        return frame_paths

    def upload_user_media(self, media_paths: List[str]):
        """Upload user media files to Gemini"""
        for path in media_paths:
            file_path = Path(path)
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    uploaded_file = genai.upload_file(str(file_path))
                    self.document_dict['user_data'][file_path.name] = uploaded_file
                except Exception as e:
                    print(f"Error loading user media {file_path.name}: {str(e)}")
            elif file_path.suffix.lower() in {'.mp4', '.mov', '.avi'}:
                self._upload_video(str(file_path))

    def _upload_video(self, video_path: str):
        """Upload video file to Gemini"""
        print("Uploading video file...")
        video_file = genai.upload_file(path=video_path)
        print(f"Completed upload: {video_file.uri}")
        
        while video_file.state.name == "PROCESSING":
            print('Waiting for video to be processed.')
            time.sleep(10)
            video_file = genai.get_file(video_file.name)
        
        if video_file.state.name == "FAILED":
            raise ValueError(video_file.state.name)
        
        print(f'Video processing complete: {video_file.uri}')
        self.document_dict['user_data'][Path(video_path).name] = video_file

    def generate_report(self) -> Dict:
        """Generate inspection report based on uploaded media"""
        prompt = """
You have been supplied with a set of building standards and manufacturer specifications to evaluate the photos and videos against.
Please be specific about any violations of building codes or manufacturer specifications found in the documentation.

Analyze the uploaded photos and videos of the building and generate a detailed inspection report in JSON format.
Be exhaustive in your inspection and cover all aspects of the building shown in the media.

The response should be a valid JSON object with the following structure:

{
  "detailedInspection": [
    {
      "area": "string",
      "mediaReference": "string",
      "timestamp": "string",
      "condition": "string",
      "complianceStatus": "string",
      "issuesFound": ["string"],
      "referenceDoc": "string",
      "referenceSection": "string",
      "recommendation": "string"
    }
  ],
  "executiveSummary": {
    "overallCondition": "string",
    "criticalIssues": ["string"],
    "recommendedActions": ["string"]
  },
  "maintenanceNotes": {
    "recurringIssues": ["string"],
    "preventiveRecommendations": ["string"],
    "maintenanceSchedule": [
      {
        "frequency": "string",
        "tasks": ["string"]
      }
    ],
    "costConsiderations": ["string"]
  }
}

Ensure the response is a valid JSON object that can be parsed.
"""
        
        content = [{'text': prompt}]
        
        # Add user media
        content.append({'text': 'User provided media:'})
        for name, doc in self.document_dict['user_data'].items():
            content.append({'text': f"User Document: {name}"})
            content.append(doc)
        
        # Start chat session
        chat_session = self.model.start_chat(history=[{"role": "user", "parts": content}])
        
        # Get response
        response = chat_session.send_message(
            "Please generate a detailed building report. "
            "Please provide a detailed answer with elaboration on the report and reference material."
        )
        
        return json.loads(response.text)

    def generate_word_report(self, report_data: Dict, output_path: str = "inspection_report.docx"):
        """Generate a professional Word document report with images and formatting"""
        doc = Document()
        
        # Add report header
        self._add_report_header(doc)
        
        # Add executive summary
        self._add_executive_summary(doc, report_data)
        
        # Add detailed findings with images
        self._add_detailed_findings(doc, report_data)
        
        # Add maintenance schedule
        self._add_maintenance_schedule(doc, report_data)
        
        # Add footer with page numbers
        self._add_footer(doc)
        
        # Save the document
        doc.save(output_path)
        return output_path

    def _add_report_header(self, doc):
        """Add report title header"""
        # Add title
        title = doc.add_heading('HOME INSPECTION REPORT', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add date and prepared by
        date_para = doc.add_paragraph()
        date_para.add_run(f"Date: {datetime.datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        prepared_para = doc.add_paragraph()
        prepared_para.add_run("Prepared by: AI Home Inspection System")
        prepared_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph()  # Add spacing

    def _add_executive_summary(self, doc, report_data):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', level=1)
        
        # Overall condition
        p = doc.add_paragraph()
        p.add_run("Overall Condition: ").bold = True
        p.add_run(report_data['executiveSummary']['overallCondition'])
        
        # Critical issues
        doc.add_heading('Critical Issues', level=2)
        for issue in report_data['executiveSummary']['criticalIssues']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(issue).bold = True
            p.add_run().font.color.rgb = RGBColor(255, 0, 0)  # Red color for critical issues
        
        # Recommended actions
        doc.add_heading('Recommended Actions', level=2)
        for action in report_data['executiveSummary']['recommendedActions']:
            doc.add_paragraph(action, style='List Bullet')
        
        doc.add_page_break()

    def _add_detailed_findings(self, doc, report_data):
        """Add detailed findings with images"""
        doc.add_heading('Detailed Inspection Findings', level=1)
        
        for finding in report_data['detailedInspection']:
            # Add finding header
            doc.add_heading(f"{finding['area']} - {finding['condition']}", level=2)
            
            # Try to add image if available
            if finding.get('mediaReference'):
                media_ref = finding['mediaReference']
                if media_ref.startswith('frame_'):
                    frame_path = os.path.join("extracted_frames", media_ref)
                    if os.path.exists(frame_path):
                        try:
                            # Add image with caption
                            doc.add_picture(frame_path, width=Inches(4.0))
                            caption = doc.add_paragraph(f"Figure: {finding['area']} at {finding.get('timestamp', 'unknown time')}")
                            caption.style = 'Caption'
                            doc.add_paragraph()  # Add spacing after image
                        except Exception as e:
                            print(f"Error adding image to Word doc: {e}")
            
            # Add compliance status with color
            p = doc.add_paragraph()
            p.add_run("Compliance Status: ").bold = True
            status_run = p.add_run(finding['complianceStatus'])
            if finding['complianceStatus'] == 'Non-compliant':
                status_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            else:
                status_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            
            # Add issues found
            if finding.get('issuesFound'):
                doc.add_heading('Issues Found', level=3)
                for issue in finding['issuesFound']:
                    doc.add_paragraph(issue, style='List Bullet')
            
            # Add reference
            if finding.get('referenceDoc') and finding.get('referenceSection'):
                p = doc.add_paragraph()
                p.add_run("Standard Reference: ").bold = True
                p.add_run(f"{finding['referenceDoc']} - {finding['referenceSection']}")
            
            # Add recommendation
            if finding.get('recommendation'):
                p = doc.add_paragraph()
                p.add_run("Recommendation: ").bold = True
                p.add_run(finding['recommendation'])
            
            doc.add_paragraph()  # Add space between findings

    def _add_maintenance_schedule(self, doc, report_data):
        """Add maintenance schedule section"""
        doc.add_heading('Maintenance Schedule', level=1)
        
        # Maintenance tasks
        for schedule in report_data['maintenanceNotes']['maintenanceSchedule']:
            doc.add_heading(f"{schedule['frequency']} Tasks", level=2)
            for task in schedule['tasks']:
                doc.add_paragraph(task, style='List Bullet')
        
        # Cost considerations
        if report_data['maintenanceNotes'].get('costConsiderations'):
            doc.add_heading('Cost Considerations', level=2)
            for cost in report_data['maintenanceNotes']['costConsiderations']:
                doc.add_paragraph(cost, style='List Bullet')

    def _add_footer(self, doc):
        """Add footer with page numbers to all sections"""
        section = doc.sections[0]
        footer = section.footer
        
        # Add page number
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add page number field
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run._r.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)
